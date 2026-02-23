"""
Document generation classes for PDF and DOCX export.
"""

import os
import io
import qrcode
import logging
from datetime import datetime
from io import BytesIO
from typing import Dict, Any, List

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import black, blue, grey
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class PDFGenerator:
    """Generate PDF documents from CV content."""

    def __init__(self, template_config: Dict[str, Any] = None):
        self.template_config = template_config or {}
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom paragraph styles."""
        # Header style
        self.styles.add(ParagraphStyle(
            name='CustomHeader',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            textColor=black,
            alignment=TA_CENTER
        ))

        # Section header style
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=6,
            textColor=black,
            alignment=TA_LEFT
        ))

        # Professional summary style
        self.styles.add(ParagraphStyle(
            name='Summary',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_LEFT
        ))

        # Job title style
        self.styles.add(ParagraphStyle(
            name='JobTitle',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceBefore=6,
            spaceAfter=3,
            textColor=black,
            alignment=TA_LEFT
        ))

        # Company/date style
        self.styles.add(ParagraphStyle(
            name='CompanyDate',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            textColor=grey,
            alignment=TA_LEFT
        ))

    def generate_cv(self, content: Dict[str, Any], options: Dict[str, Any] = None) -> bytes:
        """Generate CV PDF from content."""
        if options is None:
            options = {}

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )

        story = []

        # Header section
        story.extend(self._create_header(content.get('personal_info', {})))

        # Professional summary
        if content.get('professional_summary') and options.get('include_professional_summary', True):
            story.extend(self._create_section(
                "Professional Summary",
                content['professional_summary'],
                style='Summary'
            ))

        # Key skills
        if content.get('key_skills') and options.get('include_skills', True):
            story.extend(self._create_skills_section(content['key_skills']))

        # Experience
        if content.get('experience') and options.get('include_experience', True):
            story.extend(self._create_experience_section(
                content['experience'],
                options.get('include_evidence', False)
            ))

        # Projects
        if content.get('projects') and options.get('include_projects', True):
            story.extend(self._create_projects_section(
                content['projects'],
                options.get('include_evidence', False)
            ))

        # Education
        if content.get('education') and options.get('include_education', True):
            story.extend(self._create_education_section(content['education']))

        # Certifications
        if content.get('certifications') and options.get('include_certifications', True):
            story.extend(self._create_certifications_section(content['certifications']))

        # Evidence section (if requested)
        if options.get('include_evidence', False):
            evidence_links = self._collect_evidence_links(content)
            if evidence_links:
                story.extend(self._create_evidence_section(evidence_links, options))

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _create_header(self, personal_info: Dict[str, Any]) -> List:
        """Create header section with personal information."""
        story = []

        # Name
        name = personal_info.get('name', 'Your Name')
        story.append(Paragraph(name, self.styles['CustomHeader']))

        # Contact information
        contact_parts = []
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_parts.append(personal_info['location'])

        if contact_parts:
            contact_text = ' | '.join(contact_parts)
            story.append(Paragraph(contact_text, self.styles['Normal']))

        # Links
        links = []
        if personal_info.get('linkedin_url'):
            links.append(f'<link href="{personal_info["linkedin_url"]}">LinkedIn</link>')
        if personal_info.get('github_url'):
            links.append(f'<link href="{personal_info["github_url"]}">GitHub</link>')
        if personal_info.get('website_url'):
            links.append(f'<link href="{personal_info["website_url"]}">Portfolio</link>')

        if links:
            links_text = ' | '.join(links)
            story.append(Paragraph(links_text, self.styles['Normal']))

        story.append(Spacer(1, 12))
        story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=grey))
        story.append(Spacer(1, 12))

        return story

    def _create_section(self, title: str, content: str, style: str = 'Normal') -> List:
        """Create a section with title and content."""
        story = []
        story.append(Paragraph(title, self.styles['SectionHeader']))
        story.append(Paragraph(content, self.styles[style]))
        story.append(Spacer(1, 12))
        return story

    def _create_skills_section(self, skills: List[str]) -> List:
        """Create skills section."""
        story = []
        story.append(Paragraph("Key Skills", self.styles['SectionHeader']))

        # Format skills as comma-separated string
        skills_text = ', '.join(skills)
        story.append(Paragraph(skills_text, self.styles['Normal']))
        story.append(Spacer(1, 12))

        return story

    def _create_experience_section(self, experience: List[Dict[str, Any]], include_evidence: bool = False) -> List:
        """Create experience section."""
        story = []
        story.append(Paragraph("Professional Experience", self.styles['SectionHeader']))

        for exp in experience:
            # Job title and company
            title_text = f"<b>{exp.get('title', 'Position')}</b>"
            story.append(Paragraph(title_text, self.styles['JobTitle']))

            company_date = f"{exp.get('organization', 'Company')} | {exp.get('duration', 'Duration')}"
            story.append(Paragraph(company_date, self.styles['CompanyDate']))

            # Achievements
            for achievement in exp.get('achievements', []):
                story.append(Paragraph(f"• {achievement}", self.styles['Normal']))

            # Technologies
            if exp.get('technologies_used'):
                tech_text = f"<i>Technologies: {', '.join(exp['technologies_used'])}</i>"
                story.append(Paragraph(tech_text, self.styles['Normal']))

            story.append(Spacer(1, 8))

        return story

    def _create_projects_section(self, projects: List[Dict[str, Any]], include_evidence: bool = False) -> List:
        """Create projects section."""
        story = []
        story.append(Paragraph("Key Projects", self.styles['SectionHeader']))

        for project in projects:
            # Project name
            name_text = f"<b>{project.get('name', 'Project')}</b>"
            story.append(Paragraph(name_text, self.styles['JobTitle']))

            # Description
            description = project.get('description', '')
            if description:
                story.append(Paragraph(description, self.styles['Normal']))

            # Technologies
            if project.get('technologies'):
                tech_text = f"<i>Technologies: {', '.join(project['technologies'])}</i>"
                story.append(Paragraph(tech_text, self.styles['Normal']))

            # Impact metrics
            if project.get('impact_metrics'):
                impact_text = f"<i>Impact: {project['impact_metrics']}</i>"
                story.append(Paragraph(impact_text, self.styles['Normal']))

            story.append(Spacer(1, 8))

        return story

    def _create_education_section(self, education: List[Dict[str, Any]]) -> List:
        """Create education section."""
        story = []
        if not education:
            return story

        story.append(Paragraph("Education", self.styles['SectionHeader']))

        for edu in education:
            degree_text = f"<b>{edu.get('degree', 'Degree')}</b>"
            story.append(Paragraph(degree_text, self.styles['JobTitle']))

            institution_year = f"{edu.get('institution', 'Institution')} | {edu.get('year', 'Year')}"
            story.append(Paragraph(institution_year, self.styles['CompanyDate']))

            if edu.get('details'):
                story.append(Paragraph(edu['details'], self.styles['Normal']))

            story.append(Spacer(1, 8))

        return story

    def _create_certifications_section(self, certifications: List[Dict[str, Any]]) -> List:
        """Create certifications section."""
        story = []
        if not certifications:
            return story

        story.append(Paragraph("Certifications", self.styles['SectionHeader']))

        for cert in certifications:
            cert_text = f"<b>{cert.get('name', 'Certification')}</b>"
            story.append(Paragraph(cert_text, self.styles['JobTitle']))

            issuer_date = f"{cert.get('issuer', 'Issuer')} | {cert.get('date', 'Date')}"
            story.append(Paragraph(issuer_date, self.styles['CompanyDate']))

            story.append(Spacer(1, 8))

        return story

    def _collect_evidence_links(self, content: Dict[str, Any]) -> List[str]:
        """Collect all evidence links from content."""
        links = []

        # From experience
        for exp in content.get('experience', []):
            links.extend(exp.get('evidence_references', []))

        # From projects
        for project in content.get('projects', []):
            if project.get('evidence_url'):
                links.append(project['evidence_url'])

        return list(set(links))  # Remove duplicates

    def _create_evidence_section(self, evidence_links: List[str], options: Dict[str, Any]) -> List:
        """Create evidence section with links."""
        story = []
        if not evidence_links:
            return story

        story.append(Paragraph("Supporting Evidence", self.styles['SectionHeader']))

        evidence_format = options.get('evidence_format', 'hyperlinks')

        if evidence_format == 'qr_codes':
            # Generate QR codes for links (simplified for example)
            for i, link in enumerate(evidence_links, 1):
                story.append(Paragraph(f"{i}. {link}", self.styles['Normal']))
        else:
            # Standard hyperlinks or footnotes
            for i, link in enumerate(evidence_links, 1):
                link_text = f'{i}. <link href="{link}">{link}</link>'
                story.append(Paragraph(link_text, self.styles['Normal']))

        return story


class DOCXGenerator:
    """Generate DOCX documents from CV content."""

    def __init__(self, template_path: str = None):
        self.template_path = template_path

    def generate_cv(self, content: Dict[str, Any], options: Dict[str, Any] = None) -> bytes:
        """Generate CV DOCX from content."""
        if options is None:
            options = {}

        doc = Document()

        # Set document properties
        personal_info = content.get('personal_info', {})
        doc.core_properties.title = f"CV - {personal_info.get('name', 'Your Name')}"
        doc.core_properties.author = personal_info.get('name', 'Your Name')

        # Header
        self._add_header(doc, personal_info)

        # Professional summary
        if content.get('professional_summary') and options.get('include_professional_summary', True):
            self._add_section(doc, "Professional Summary", content['professional_summary'])

        # Skills
        if content.get('key_skills') and options.get('include_skills', True):
            self._add_skills_section(doc, content['key_skills'])

        # Experience
        if content.get('experience') and options.get('include_experience', True):
            self._add_experience_section(doc, content['experience'], options)

        # Projects
        if content.get('projects') and options.get('include_projects', True):
            self._add_projects_section(doc, content['projects'], options)

        # Education
        if content.get('education') and options.get('include_education', True):
            self._add_education_section(doc, content['education'])

        # Certifications
        if content.get('certifications') and options.get('include_certifications', True):
            self._add_certifications_section(doc, content['certifications'])

        return self._save_to_buffer(doc)

    def _add_header(self, doc: Document, personal_info: Dict[str, Any]):
        """Add header with personal information."""
        # Name
        name_para = doc.add_heading(personal_info.get('name', 'Your Name'), level=1)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_parts = []
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_parts.append(personal_info['location'])

        if contact_parts:
            contact_para = doc.add_paragraph(' | '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Links
        links = []
        if personal_info.get('linkedin_url'):
            links.append(f"LinkedIn: {personal_info['linkedin_url']}")
        if personal_info.get('github_url'):
            links.append(f"GitHub: {personal_info['github_url']}")
        if personal_info.get('website_url'):
            links.append(f"Portfolio: {personal_info['website_url']}")

        if links:
            links_para = doc.add_paragraph(' | '.join(links))
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_section(self, doc: Document, title: str, content: str):
        """Add a section with title and content."""
        doc.add_heading(title, level=2)
        doc.add_paragraph(content)

    def _add_skills_section(self, doc: Document, skills: List[str]):
        """Add skills section."""
        doc.add_heading("Key Skills", level=2)
        skills_text = ', '.join(skills)
        doc.add_paragraph(skills_text)

    def _add_experience_section(self, doc: Document, experience: List[Dict[str, Any]], options: Dict[str, Any]):
        """Add experience section."""
        doc.add_heading("Professional Experience", level=2)

        for exp in experience:
            # Job title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(exp.get('title', 'Position'))
            title_run.bold = True

            # Company and duration
            company_para = doc.add_paragraph(f"{exp.get('organization', 'Company')} | {exp.get('duration', 'Duration')}")

            # Achievements
            for achievement in exp.get('achievements', []):
                doc.add_paragraph(f"• {achievement}")

            # Technologies
            if exp.get('technologies_used'):
                tech_para = doc.add_paragraph()
                tech_para.add_run("Technologies: ").italic = True
                tech_para.add_run(', '.join(exp['technologies_used']))

    def _add_projects_section(self, doc: Document, projects: List[Dict[str, Any]], options: Dict[str, Any]):
        """Add projects section."""
        doc.add_heading("Key Projects", level=2)

        for project in projects:
            # Project name
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(project.get('name', 'Project'))
            name_run.bold = True

            # Description
            if project.get('description'):
                doc.add_paragraph(project['description'])

            # Technologies
            if project.get('technologies'):
                tech_para = doc.add_paragraph()
                tech_para.add_run("Technologies: ").italic = True
                tech_para.add_run(', '.join(project['technologies']))

            # Impact
            if project.get('impact_metrics'):
                impact_para = doc.add_paragraph()
                impact_para.add_run("Impact: ").italic = True
                impact_para.add_run(project['impact_metrics'])

    def _add_education_section(self, doc: Document, education: List[Dict[str, Any]]):
        """Add education section."""
        if not education:
            return

        doc.add_heading("Education", level=2)

        for edu in education:
            # Degree
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(edu.get('degree', 'Degree'))
            degree_run.bold = True

            # Institution and year
            doc.add_paragraph(f"{edu.get('institution', 'Institution')} | {edu.get('year', 'Year')}")

            # Details
            if edu.get('details'):
                doc.add_paragraph(edu['details'])

    def _add_certifications_section(self, doc: Document, certifications: List[Dict[str, Any]]):
        """Add certifications section."""
        if not certifications:
            return

        doc.add_heading("Certifications", level=2)

        for cert in certifications:
            # Certification name
            cert_para = doc.add_paragraph()
            cert_run = cert_para.add_run(cert.get('name', 'Certification'))
            cert_run.bold = True

            # Issuer and date
            doc.add_paragraph(f"{cert.get('issuer', 'Issuer')} | {cert.get('date', 'Date')}")

    def _save_to_buffer(self, doc: Document) -> bytes:
        """Save document to bytes buffer."""
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()